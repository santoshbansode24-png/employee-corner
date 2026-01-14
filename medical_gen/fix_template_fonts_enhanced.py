"""
Enhanced script to fix both English and Marathi fonts in Word template
Ensures proper font embedding for PDF generation
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_run_font(run, font_name):
    """Set font for a run with proper XML attributes for PDF embedding"""
    run.font.name = font_name
    
    # Set font for different character sets
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    
    # ASCII characters (English)
    rFonts.set(qn('w:ascii'), font_name)
    # High ANSI characters
    rFonts.set(qn('w:hAnsi'), font_name)
    # East Asian characters
    rFonts.set(qn('w:eastAsia'), font_name)
    # Complex script characters (Devanagari)
    rFonts.set(qn('w:cs'), font_name)

def update_template_fonts_enhanced(template_path):
    """Update template with proper font settings for both English and Marathi"""
    
    print(f"Opening template: {template_path}")
    doc = Document(template_path)
    
    # Update Normal style
    print("Updating Normal style...")
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'  # Use Arial for better compatibility
    
    # Set font for all character sets in Normal style
    rPr = normal_style._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rFonts.set(qn('w:eastAsia'), 'Noto Sans Devanagari')
    rFonts.set(qn('w:cs'), 'Noto Sans Devanagari')
    
    # Update all paragraphs
    print("Updating paragraph fonts...")
    for paragraph in doc.paragraphs:
        # Check if paragraph contains Devanagari characters
        text = paragraph.text
        has_devanagari = any('\u0900' <= char <= '\u097F' for char in text)
        
        for run in paragraph.runs:
            run_text = run.text
            run_has_devanagari = any('\u0900' <= char <= '\u097F' for char in run_text)
            
            if run_has_devanagari:
                # Use Devanagari font for Marathi text
                set_run_font(run, 'Noto Sans Devanagari')
            else:
                # Use Arial for English text
                set_run_font(run, 'Arial')
    
    # Update all tables
    print("Updating table fonts...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    has_devanagari = any('\u0900' <= char <= '\u097F' for char in text)
                    
                    for run in paragraph.runs:
                        run_text = run.text
                        run_has_devanagari = any('\u0900' <= char <= '\u097F' for char in run_text)
                        
                        if run_has_devanagari:
                            set_run_font(run, 'Noto Sans Devanagari')
                        else:
                            set_run_font(run, 'Arial')
    
    # Create backup
    backup_path = template_path.replace('.docx', '_backup_enhanced.docx')
    print(f"Creating backup: {backup_path}")
    
    if os.path.exists(template_path):
        import shutil
        shutil.copy(template_path, backup_path)
    
    # Save updated template
    print(f"Saving updated template: {template_path}")
    doc.save(template_path)
    print("✅ Template updated successfully with enhanced font settings!")
    
    return True

if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(script_dir, "template_fixed.docx")
    
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        template_path = os.path.join(script_dir, "template.docx")
        print(f"Trying alternate template: {template_path}")
    
    if os.path.exists(template_path):
        update_template_fonts_enhanced(template_path)
    else:
        print("❌ No template found!")
