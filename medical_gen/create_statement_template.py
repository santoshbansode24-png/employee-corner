from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_template():
    doc = Document()
    
    # Title
    p = doc.add_paragraph('Statement showing the details of medicines purchased for the treatment of ')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('{{ patient_name }}')
    run.bold = True
    
    # Details Paragraph
    p2 = doc.add_paragraph()
    p2.add_run('Name of Employee: {{ emp_name }}\n')
    p2.add_run('Designation: {{ designation }}\n')
    p2.add_run('Office: {{ office_name }}')
    
    # Table
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ['Sr. No.', 'Bill No.', 'Date', 'Name of Chemist/Medical', 'Amount (Rs.)']
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    # Data Row (Template for docxtpl)
    row_cells = table.rows[1].cells
    
    # Using docxtpl tags
    row_cells[0].text = '{{ item.sr }}'
    row_cells[1].text = '{{ item.bill_no }}'
    row_cells[2].text = '{{ item.date }}'
    row_cells[3].text = '{{ item.chemist }}'
    row_cells[4].text = '{{ item.amount }}'
    
    # Add the Jinja2 loop tags AROUND the row cells in the XML is tricky with just python-docx,
    # usually we put {% tr for item in items %} in the first cell and {% endtr %} in the last cell for docxtpl.
    
    row_cells[0].text = '{% tr for item in items %}{{ item.sr }}'
    row_cells[4].text = '{{ item.amount }}{% endtr %}'
    
    # Total Row
    total_row = table.add_row().cells
    total_row[3].text = 'Total'
    total_row[3].paragraphs[0].runs[0].bold = True
    total_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row[4].text = '{{ total_amount }}'
    
    # Signature Area
    doc.add_paragraph('\n\n')
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.width = Inches(6)
    sig_cells = sig_table.rows[0].cells
    
    sig_cells[0].text = 'Signature of Applicant'
    sig_cells[1].text = 'Signature of Controlling Officer'
    sig_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.save('medical_gen/statement_template.docx')
    print("Template created successfully.")

if __name__ == "__main__":
    create_template()
