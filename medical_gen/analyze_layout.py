from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def analyze_template_layout():
    """Analyze the current template layout"""
    
    if not os.path.exists("template.docx"):
        print("Error: template.docx not found!")
        return
    
    doc = Document("template.docx")
    
    print("=" * 60)
    print("TEMPLATE LAYOUT ANALYSIS")
    print("=" * 60)
    
    # Page setup
    section = doc.sections[0]
    print(f"\n📄 Page Setup:")
    print(f"   Page Width: {section.page_width.inches:.2f} inches")
    print(f"   Page Height: {section.page_height.inches:.2f} inches")
    print(f"   Top Margin: {section.top_margin.inches:.2f} inches")
    print(f"   Bottom Margin: {section.bottom_margin.inches:.2f} inches")
    print(f"   Left Margin: {section.left_margin.inches:.2f} inches")
    print(f"   Right Margin: {section.right_margin.inches:.2f} inches")
    
    # Count elements
    print(f"\n📊 Document Structure:")
    print(f"   Total Paragraphs: {len(doc.paragraphs)}")
    print(f"   Total Tables: {len(doc.tables)}")
    
    # Analyze tables
    print(f"\n📋 Tables Analysis:")
    for i, table in enumerate(doc.tables, 1):
        print(f"\n   Table {i}:")
        print(f"      Rows: {len(table.rows)}")
        print(f"      Columns: {len(table.columns)}")
        
        # Check for empty cells
        empty_cells = 0
        total_cells = 0
        for row in table.rows:
            for cell in row.cells:
                total_cells += 1
                if not cell.text.strip():
                    empty_cells += 1
        
        print(f"      Total Cells: {total_cells}")
        print(f"      Empty Cells: {empty_cells}")
        print(f"      Fill Rate: {((total_cells-empty_cells)/total_cells*100):.1f}%")
    
    print("\n" + "=" * 60)
    print("💡 RECOMMENDATIONS:")
    print("=" * 60)
    print("""
1. Open template.docx in Microsoft Word
2. Adjust margins: Layout → Margins → Custom Margins
   - Recommended: Top/Bottom: 0.5", Left/Right: 0.75"
   
3. For tables:
   - Right-click table → Table Properties
   - Set preferred width to 100%
   - Adjust row heights if needed
   
4. For spacing:
   - Select all text (Ctrl+A)
   - Paragraph → Line spacing: Single
   - Remove space before/after paragraphs
   
5. For page breaks:
   - Insert → Page Break where needed
   - Or use Ctrl+Enter
    """)

if __name__ == "__main__":
    analyze_template_layout()
