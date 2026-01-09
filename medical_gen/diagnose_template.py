from docx import Document
import re

def list_template_variables(docx_path):
    try:
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
                    
        content = "\n".join(full_text)
        print("--- Content Dump ---")
        # print(content) 
        
        # Find potential jinja2 tags
        matches = re.findall(r'\{\{.*?\}\}', content)
        print(f"\nFound {len(matches)} valid-looking tags:")
        for m in matches:
            print(m)
            
        # Find half tags
        opens = re.findall(r'\{\{', content)
        closes = re.findall(r'\}\}', content)
        print(f"\nTotal '{{{{': {len(opens)}")
        print(f"Total '}}}}': {len(closes)}")
        
        if len(opens) != len(closes):
            print("MISMATCH DETECTED!")
            
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    list_template_variables("template.docx")
