from docx import Document
import os

def fix_table_dates():
    script_dir = os.path.dirname(__file__)
    files_to_fix = ["template_fixed.docx", "template.docx"]
    
    for filename in files_to_fix:
        template_path = os.path.join(script_dir, filename)
        
        if not os.path.exists(template_path):
            print(f"Error: {template_path} not found.")
            continue

        print(f"Fixing dates in: {template_path}")
        doc = Document(template_path)
        
        # Target Table 1 (Hospital Stay Table)
        if len(doc.tables) < 2:
            print(f"Error: Not enough tables found in {filename}.")
            continue
        
        table = doc.tables[1]
        
        # Changes to apply: (Row Index, Column Index, Correct Variable)
        # Indices are 0-based.
        # Column 2 is the 'Dates' column.
        
        changes = [
            (1, 2, "{{gw_dates}}"),      # General Ward
            (2, 2, "{{semi_dates}}"),    # Semi-Private
            (3, 2, "{{pvt_dates}}"),     # Private Room
            (6, 2, "{{icu_dates}}")      # ICU
        ]
        
        for r_idx, c_idx, new_text in changes:
            try:
                cell = table.rows[r_idx].cells[c_idx]
                print(f"Row {r_idx} Col {c_idx} Original: '{cell.text}'")
                
                # Use specific replacement logic to be safe
                if "{{" in cell.text:
                    # Fully replace content to ensure no duplication
                    cell.text = new_text
                    print(f"  -> Updated to: '{cell.text}'")
                else:
                    print("  -> No variable found, checking if correct...")
                    
            except IndexError:
                print(f"Error: Row {r_idx} Col {c_idx} out of bounds.")
                
        doc.save(template_path)
        print(f"✅ {filename} successfully updated!\n")

if __name__ == "__main__":
    fix_table_dates()
