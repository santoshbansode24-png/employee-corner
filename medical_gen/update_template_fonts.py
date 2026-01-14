"""
Script to update the Word template to use Noto Sans Devanagari font
This ensures Marathi text renders correctly in the generated PDF
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

def update_template_fonts(template_path):
    """Update all fonts in the template to use Noto Sans Devanagari"""
    
    print(f"Opening template: {template_path}")
    doc = Document(template_path)
    
    # Update default font for the document
    print("Updating document default font...")
    doc.styles['Normal'].font.name = 'Noto Sans Devanagari'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Devanagari')
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:cs'), 'Noto Sans Devanagari')
    
    # Update all paragraphs
    print("Updating paragraph fonts...")
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Noto Sans Devanagari'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Devanagari')
            run._element.rPr.rFonts.set(qn('w:cs'), 'Noto Sans Devanagari')
    
    # Update all tables
    print("Updating table fonts...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Noto Sans Devanagari'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Devanagari')
                        run._element.rPr.rFonts.set(qn('w:cs'), 'Noto Sans Devanagari')
    
    # Save the updated template
    backup_path = template_path.replace('.docx', '_backup.docx')
    print(f"Creating backup: {backup_path}")
    
    if os.path.exists(template_path):
        import shutil
        shutil.copy(template_path, backup_path)
    
    print(f"Saving updated template: {template_path}")
    doc.save(template_path)
    print("✅ Template updated successfully!")
    
    return True

if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(script_dir, "template_fixed.docx")
    
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        template_path = os.path.join(script_dir, "template.docx")
        print(f"Trying alternate template: {template_path}")
    
    if os.path.exists(template_path):
        update_template_fonts(template_path)
    else:
        print("❌ No template found!")
