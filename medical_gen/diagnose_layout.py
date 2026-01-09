
from docx import Document
from docx.shared import Inches

def diagnose():
    filename = 'template_fixed.docx'
    doc = Document(filename)
    
    print(f"--- Document: {filename} ---")
    
    # Check sections (margins)
    for i, section in enumerate(doc.sections):
        print(f"Section {i+1}:")
        print(f"  Page Width: {section.page_width.inches} inches")
        print(f"  Left Margin: {section.left_margin.inches} inches")
        print(f"  Right Margin: {section.right_margin.inches} inches")
        printable_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
        print(f"  Printable Width: {printable_width} inches")

    # Check tables
    print(f"\nFound {len(doc.tables)} tables.")
    for i, table in enumerate(doc.tables):
        print(f"Table {i+1}:")
        print(f"  Autofit: {table.autofit}")
        # Estimating width by summing cell widths of first row if available
        if len(table.rows) > 0:
            row = table.rows[0]
            total_width = 0
            for cell in row.cells:
                if cell.width:
                    total_width += cell.width.inches
            print(f"  Approx Width (Row 1): {total_width} inches")
        else:
            print("  (Empty table)")

if __name__ == "__main__":
    diagnose()
