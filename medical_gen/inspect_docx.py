import re
from docx import Document
import sys

def get_docx_tags(docx_path):
    try:
        doc = Document(docx_path)
        text = ""
        for p in doc.paragraphs:
            text += p.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        # Find tags with some context
        import re
        # Find all text and identify lines with tags
        lines = text.split('\n')
        print(f"Tags with context in {docx_path}:")
        for line in lines:
            if "{{" in line:
                print(f"Line: {line.strip()}")
                
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        get_docx_tags(sys.argv[1])
    else:
        print("Usage: python inspect_docx.py <path_to_docx>")
